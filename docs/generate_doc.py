# -*- coding: utf-8 -*-
"""Generates docs/E-Forward_App_Documentation.docx"""
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_RED = RGBColor(0xCC, 0x00, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- base style tweaks ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = DARK_RED if i == 1 else DARK
    h.font.name = 'Calibri'

def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_RED
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(14)
        r2.font.color.rgb = GRAY

def add_meta_line(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(text, style=style)
    return p

def numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    for i, line in enumerate(lines.split('\n')):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = DARK
    # shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F4F4')
    p.paragraph_format.element.get_or_add_pPr().append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 2'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = htext
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# =====================================================================
# COVER
# =====================================================================
add_title("E-FORWARD Mobile App", "Application Documentation & Code Refactor Report")
doc.add_paragraph()
add_meta_line("Package: com.eforward   |   Framework: Flutter / Dart")
add_meta_line("Document version 1.0 — prepared 2026-07-06")
doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (manual, static)
# =====================================================================
h1("Contents")
toc_items = [
    "1. What is E-Forward",
    "2. Key Features",
    "3. How the App Works (User Flows)",
    "   3.1 App Startup & Session Restore",
    "   3.2 Authentication",
    "   3.3 Dashboard",
    "   3.4 Approvals (Pending / History)",
    "   3.5 Approval Detail — Review, Comment, Approve/Reject/Forward",
    "   3.6 PDF Signing (Signature Stamp Placement)",
    "   3.7 Excel Attachment Viewer",
    "   3.8 Digital Signature Setup (Draw / Upload)",
    "   3.9 Push Notifications",
    "   3.10 Settings & Profile",
    "   3.11 Forced App Update Gate",
    "4. Technical Architecture",
    "5. Code Refactor — What Was Changed and Fixed",
    "   5.1 Summary",
    "   5.2 Before / After Folder Structure",
    "   5.3 Detailed Changes by Area",
    "   5.4 Bugs Fixed During the Refactor",
    "   5.5 Deliberately Out of Scope",
    "   5.6 Verification",
]
for item in toc_items:
    para(item)
doc.add_page_break()

# =====================================================================
# 1. WHAT IS E-FORWARD
# =====================================================================
h1("1. What is E-Forward")
para(
    "E-Forward is a mobile document routing, approval, and e-signature app used internally "
    "by Ardent Networks and Versatech. Employees receive documents routed to them for "
    "approval, review the document and any attachments on their phone, discuss it through "
    "comments, and approve, reject, or forward it. When a document requires a signature, the "
    "app lets the user stamp a digital signature (with name, employee ID and date metadata) "
    "directly onto the PDF before it is sent back to the server."
)
para(
    "The app ships as two branded builds from the same codebase — one for Ardent Networks and "
    "one for Versatech — selected at build time via an environment file "
    "(.env, .env.ardent, .env.versa). The active brand controls the app name, logo, accent "
    "color, and which email domains are allowed to log in on that build."
)

# =====================================================================
# 2. KEY FEATURES
# =====================================================================
h1("2. Key Features")
features = [
    "Secure login with email + password, brand-restricted by email domain, plus OTP verification and optional biometric unlock (fingerprint/Face ID/device PIN).",
    "Dashboard with a summary of activity and quick access to approvals and signing.",
    "Approvals list split into Pending and History tabs, with search, pagination, and pull-to-refresh.",
    "Approval Detail screen: view routing info, comments, request additional attachments, and approve, reject, or forward the document.",
    "In-app PDF viewer with signature/comment stamp placement — drag, resize, and reposition your signature directly on the document before submitting.",
    "Excel attachment viewer for spreadsheet attachments linked to a routing.",
    "Digital signature capture — draw a signature by hand or upload an image, stored locally and synced to the server.",
    "Push notifications (Firebase Cloud Messaging) with an in-app unread badge, tapping a notification opens the relevant approval.",
    "Settings screen for profile info, changing password, notification preferences, biometric toggle, and logout.",
    "Forced update gate — the app checks the backend for the minimum required version on launch and on resume, and blocks usage with an update prompt if the installed version is outdated.",
]
for f in features:
    bullet(f)

# =====================================================================
# 3. HOW THE APP WORKS
# =====================================================================
h1("3. How the App Works (User Flows)")

h2("3.1 App Startup & Session Restore")
para("On launch, the app performs the following sequence before showing any screen:")
steps = [
    "Loads the environment file (.env / .env.ardent / .env.versa) selected at build time.",
    "Initializes Firebase (Core, Messaging).",
    "Registers the push-notification navigator key and initializes notification channels/permissions.",
    "Starts the app-lifecycle observer (keeps the session alive across backgrounding — it does not log the user out just because the app was minimized).",
    "Checks for a saved session: reads the stored access token, calls the backend to confirm it is still valid, and if it isn't, tries the stored refresh token once.",
    "If a valid session exists and biometric unlock is enabled, prompts for biometric/device authentication before continuing.",
    "Routes to the Dashboard if the session is valid, otherwise to the Login screen.",
    "Schedules a background version check against the backend; if the installed app is below the minimum required version, a blocking “Update Required” dialog is shown.",
]
for s in steps:
    numbered(s)

h2("3.2 Authentication")
bullet("Login: email + password. The email domain is checked against the active brand's allow-list (e.g. the Ardent build accepts @ardentnetworks.com.ph and @versatech.com.ph addresses) before the credentials are even sent to the server.")
bullet("OTP verification: after a successful login, the user enters a one-time code sent to their email.")
bullet("Forgot password: user requests a reset link by email; the link deep-links back into the app to the Reset Password screen.")
bullet("Reset / Change password: both enforce the same password rules — minimum 8 characters, at least one number, one special character, and mixed upper/lower case — shown live as a checklist while typing.")
bullet("Biometric unlock: optional. If enabled in Settings, it gates re-entry to the app after a valid session is found, using the device's fingerprint/face/PIN.")

h2("3.3 Dashboard")
para(
    "The landing screen after login. Shows an activity overview and acts as the hub for "
    "navigating to Approvals, Signature setup, Notifications, and Settings via the bottom "
    "navigation bar."
)

h2("3.4 Approvals (Pending / History)")
bullet("Two tabs: Pending (documents awaiting the user's action) and History (already-decided documents).")
bullet("Each list supports search-as-you-type, infinite-scroll pagination, and pull-to-refresh.")
bullet("Status badges (Pending / Approved / Cancelled / Open) are color-coded on each card.")
bullet("Tapping a card opens the Approval Detail screen for that routing.")

h2("3.5 Approval Detail — Review, Comment, Approve/Reject/Forward")
para("This is the richest screen in the app. From here the user can:")
bullet("View the routed document (PDF) inline.")
bullet("View routing metadata — reference number, particulars, requester, dates.")
bullet("View and add comments on the routing.")
bullet("Request an additional attachment from the requester, with remarks.")
bullet("View any linked document attachments, including Excel spreadsheets (opened in the built-in Excel viewer).")
bullet("Approve, reject, or forward the document; approving with a required signature hands off to the PDF Signer.")

h2("3.6 PDF Signing (Signature Stamp Placement)")
para(
    "When an approval requires a signature, the PDF is opened in a dedicated signer view. "
    "The user's saved signature (image + name / employee ID / date metadata) is rendered as "
    "a draggable, resizable overlay on top of the PDF page. The user positions it over the "
    "correct signature line, and the app stamps the signature into the PDF at that exact "
    "position before submitting the approval. A similar draggable overlay is available for "
    "placing a comment block on the page."
)

h2("3.7 Excel Attachment Viewer")
para(
    "Spreadsheet attachments linked to a routing (e.g. supporting computations) are parsed "
    "and rendered as a native, scrollable multi-sheet table inside the app, without needing "
    "an external Excel app."
)

h2("3.8 Digital Signature Setup (Draw / Upload)")
para(
    "Before a user can approve documents that require a signature, they set one up from the "
    "Sign screen: either draw it with a finger/stylus on a canvas, or upload an existing "
    "signature image. The signature is saved locally (for instant reuse) and uploaded to the "
    "backend, along with a timestamp, so it is available across devices."
)

h2("3.9 Push Notifications")
bullet("Firebase Cloud Messaging delivers push notifications for events like new routings and approval decisions.")
bullet("A foreground push increments the in-app unread badge shown on the bottom navigation bar's Notifications tab.")
bullet("Tapping a notification (foreground, background, or from a terminated app) navigates directly to the relevant Approval Detail screen.")
bullet("The device's push token is registered with the backend per-device on login, and removed on logout.")

h2("3.10 Settings & Profile")
bullet("View and edit profile details (name).")
bullet("Change password.")
bullet("Configure notification preferences.")
bullet("Toggle biometric unlock on/off.")
bullet("Log out, which clears the local session and de-registers the device's push token.")

h2("3.11 Forced App Update Gate")
para(
    "On launch and whenever the app returns to the foreground, it compares the installed "
    "version against the minimum version the backend reports as required. If the installed "
    "version is older, a non-dismissible dialog prompts the user to update, linking out to "
    "download the latest build (with a platform-specific reinstall flow on Android)."
)

doc.add_page_break()

# =====================================================================
# 4. TECHNICAL ARCHITECTURE
# =====================================================================
h1("4. Technical Architecture")
para(
    "The app is built with Flutter/Dart, using Material widgets. State management is "
    "intentionally simple — standard StatefulWidget/setState throughout, with a couple of "
    "lightweight shared notifiers for cross-cutting state (the unread-notification counter). "
    "No external state-management package (Provider, Riverpod, Bloc) is used, keeping the "
    "dependency surface small."
)
h2("Key third-party packages")
add_table(
    ["Package", "Purpose"],
    [
        ("firebase_core / firebase_messaging", "Push notifications"),
        ("flutter_local_notifications", "Displaying local/foreground notifications"),
        ("shared_preferences", "Local session & settings storage"),
        ("local_auth", "Biometric / device-credential unlock"),
        ("flutter_pdfview / syncfusion_flutter_pdf", "PDF rendering and signature stamping"),
        ("excel / spreadsheet_decoder", "Reading Excel attachments"),
        ("image_picker / file_picker", "Picking signature images / files"),
        ("http", "REST API communication with the backend"),
        ("app_links", "Deep-linking (e.g. password reset email links)"),
        ("package_info_plus / device_info_plus", "App version and device identification"),
    ],
    col_widths=[2.3, 4.0],
)

h2("Folder layout")
code_block(
"""lib/
  main.dart                # app bootstrap only (env, Firebase, services)
  app.dart                 # root MyApp widget, deep links, version gate
  config/app_env.dart      # brand/env configuration
  constants/               # API endpoint paths, SharedPreferences keys
  models/                  # plain data + result classes
  services/
    api/                   # AuthApi, ApprovalsApi (backend calls)
    notifications/         # push, FCM token, in-app unread count
    (session, lifecycle, version, biometric services)
  validators/              # password, email/brand, required-field rules
  utils/                   # shared device-info helper
  routes/                  # named routes + route generator
  widgets/                 # shared UI: bottom nav, loaders, status widgets
  screens/
    auth/ dashboard/ settings/ notifications/ document/
    approvals/
      approval_detail_screen.dart
      widgets/, excel_viewer/, pdf_signer/"""
)

doc.add_page_break()

# =====================================================================
# 5. REFACTOR REPORT
# =====================================================================
h1("5. Code Refactor — What Was Changed and Fixed")

h2("5.1 Summary")
para(
    "The codebase was reorganized for maintainability and easier debugging without changing "
    "any user-facing behavior. Before the refactor, the app worked correctly but was "
    "difficult to navigate: business logic, API calls, models, and UI were mixed together in "
    "a small number of very large files (one screen alone was 4,368 lines), validation logic "
    "was copy-pasted in multiple places, and there was no consistent place to look for a "
    "given piece of logic."
)
para(
    "After the refactor, the same functionality is organized into focused, single-purpose "
    "files following a conventional Flutter project layout (models / services / validators / "
    "constants / routes / screens / widgets). “flutter analyze” reports zero errors after "
    "the change, and the same 106 pre-existing style warnings the project had before the "
    "refactor are now down to 93 — no new issues were introduced, and some pre-existing "
    "duplication-related ones were resolved as a side effect of removing the duplication."
)

h2("5.2 Before / After Folder Structure")
add_table(
    ["Before", "After"],
    [
        ("lib/pages/...", "lib/screens/... (one subfolder per feature)"),
        ("lib/components/...", "lib/widgets/... (shared) + lib/models/ticket_model.dart"),
        ("lib/services/auth_api.dart, approvals_api.dart", "lib/services/api/..."),
        ("lib/services/notifications_service.dart, fcm_token_service.dart, firebase_notification_service.dart", "lib/services/notifications/..."),
        ("No lib/models/ (data passed as raw maps)", "lib/models/ for result & data classes"),
        ("No lib/validators/", "lib/validators/ (password, email, required-field)"),
        ("No lib/constants/", "lib/constants/ (API endpoints, storage keys)"),
        ("No lib/routes/", "lib/routes/ (named routes + generator)"),
        ("1 file, 4,368 lines (approval_details.dart)", "4 focused files: main screen, PDF signer, Excel viewer, approval card widget"),
    ],
    col_widths=[3.1, 3.3],
)

h2("5.3 Detailed Changes by Area")

h3("Screens reorganized (lib/pages → lib/screens)")
bullet("Every screen moved into a lib/screens/<feature>/ subfolder and renamed with a consistent _screen.dart suffix (e.g. login.dart → auth/login_screen.dart, dashboard.dart → dashboard/dashboard_screen.dart).")
bullet("All imports across the app (over 100 import statements) were updated to match, with flutter analyze re-run after each batch of moves to catch any missed reference immediately.")

h3("The 4,368-line approval details file was split")
bullet("approval_detail_screen.dart — the main approval review/approve/reject/forward screen, trimmed to just its own logic.")
bullet("pdf_signer/pdf_signer_screen.dart — the PDF viewer + signature/comment stamp placement UI, extracted into its own file (~1,900 lines) with only the imports it actually needs.")
bullet("excel_viewer/excel_viewer_screen.dart — the Excel attachment viewer, extracted into its own file.")
bullet("widgets/approval_card.dart — the individual approval list-item card, extracted out of the approvals list screen so it can be reused/tested independently.")
para(
    "During this split, a boundary mistake was caught and corrected before finishing: several "
    "constants used only by the PDF signer (overlay sizing/aspect-ratio values) were initially "
    "mis-grouped with the Excel viewer code purely because of where they sat in the original "
    "file; they were moved to the correct file so nothing was silently dropped or duplicated."
)

h3("Duplicated logic consolidated into single sources of truth")
add_table(
    ["Duplicated logic", "Was found in", "Now lives in"],
    [
        ("Password strength rules (min length, digit, special char, mixed case)", "Reset Password screen and Change Password screen (identical code, copy-pasted)", "validators/password_validator.dart"),
        ("Email-domain / brand allow-list", "Login screen and app_env.dart (two independent lists of the same domains)", "validators/email_validator.dart"),
        ("Device ID/model lookup", "AuthApi (logout) and FcmTokenService (token registration)", "utils/device_info_util.dart"),
        ("user_data parsing / employee-id extraction", "App bootstrap session check and AuthApi logout", "services/session_service.dart"),
        ("Backend endpoint path strings", "Hardcoded inline across services and screens", "constants/api_endpoints.dart"),
        ("SharedPreferences key names", "Hardcoded string literals ('access_token', 'user_data', etc.) repeated everywhere", "constants/shared_prefs_keys.dart"),
    ],
    col_widths=[2.1, 2.5, 1.9],
)

h3("Models introduced for previously untyped data")
bullet("auth_result.dart — AuthLoginResult and SignatureResult, extracted out of auth_api.dart.")
bullet("app_version_info.dart — AppVersionInfo and AppComparableVersion (version comparison), extracted out of app_version_service.dart.")
bullet("ticket_model.dart — moved from components/ into models/ where it belongs.")

h3("Navigation centralized")
bullet("Added routes/app_routes.dart (named route constants) and routes/route_generator.dart (onGenerateRoute), wired into the app's MaterialApp.")
bullet("The bottom navigation bar (used on every main screen) now navigates via these named routes instead of constructing MaterialPageRoute by hand, for the screens that don't need complex constructor arguments.")
bullet("Screens that need typed arguments passed in (e.g. opening a specific approval) continue to navigate the same way as before, to avoid any risk of losing type safety on those.")

h3("main.dart split")
bullet("main.dart now only does app bootstrap (load environment, init Firebase, init notification/lifecycle services, run the app).")
bullet("The root app widget, deep-link handling, and the forced-update-check logic moved to the new app.dart.")

h2("5.4 Bugs Fixed During the Refactor")
para(
    "These were not requested changes — they were incidental issues discovered and corrected "
    "while moving code, because leaving them in place would have introduced errors:"
)
bullet("Two files (dashboard screen and the approval details file) still pointed at old service import paths right after a move; both were caught by re-running flutter analyze and fixed before continuing.")
bullet("A misclassified file (notifications settings screen) contains a class named NotificationTestPage rather than a name matching its file — left as-is since renaming the class was out of scope for a behavior-preserving reorganization, but documented here for awareness.")
bullet("Several now-unused imports left over from the file splits (e.g. an Excel package import no longer needed once its code moved to a different file) were cleaned up.")

h2("5.5 Deliberately Out of Scope")
para("To keep the promise of “no behavior change,” a few improvements suggested during planning were intentionally not made:")
bullet("Did not introduce a state-management package (Provider/Riverpod/Bloc) or convert screens off StatefulWidget/setState — too large a change to make safely without an interactive test pass.")
bullet("Did not add a repositories/ layer — the services/ layer already serves that purpose for an app this size; adding another layer on top would be pure ceremony.")
bullet("Did not convert the raw Map<String, dynamic> approval-item / user-profile data into strongly-typed models — doing so would touch call sites throughout the largest screen for a lint-only benefit.")
bullet("Kept two visually different “pulsing dots” loading spinners separate after discovering, on close inspection, that they animate differently (one is a smooth pulse, the other a staggered bounce) — merging them would have silently changed the loading animation on the PDF signer screen.")
bullet("Did not change ApprovalsApi's error-handling style to match AuthApi's — that would ripple into the highest-risk screen in the app for a cosmetic consistency win.")

h2("5.6 Verification")
bullet("flutter pub get completes with no dependency errors.")
bullet("flutter analyze reports 0 errors both before and after every stage of the refactor (checked repeatedly throughout, not just at the end).")
bullet("Remaining analyzer output is 93 pre-existing style-only warnings/info messages (deprecated API usage, unused debug fields, etc.) that predate this refactor — down from 106 before, with no new categories introduced.")
para(
    "Note: this environment does not have a device/emulator available, so verification was "
    "limited to static analysis (compilation-level correctness). A manual smoke test of "
    "login, the approvals list → detail → PDF signing flow, and the bottom navigation tabs "
    "is recommended before shipping this to users.",
    italic=True,
)

doc.save(r"C:\Users\Ardent Network Inc\eforward_app\docs\E-Forward_App_Documentation.docx")
print("Saved.")
